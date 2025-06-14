import streamlit as st
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns
import plotly.express as px
import plotly.graph_objects as go
from plotly.subplots import make_subplots
import io
import base64
from PIL import Image
import pytesseract
import PyPDF2
import docx
import json
import re
from typing import Dict, List, Any, Optional, Union
import requests
from together import Together
import warnings
import tempfile
import os

warnings.filterwarnings('ignore')

# --- Tesseract OCR Path (Uncomment and modify if pytesseract isn't found) ---
# pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe' 
# For Linux/macOS, it's usually just 'tesseract' if installed in PATH, otherwise provide full path.

# --- DataAnalystAgent Class (Your provided code) ---
class DataAnalystAgent:
    def __init__(self, api_key: str):
        """
        Initialize the Data Analyst Agent with Together.ai API
        
        Args:
            api_key (str): Together.ai API key
        """
        self.client = Together(api_key=api_key)
        self.model = "meta-llama/Llama-4-Maverick-17B-128E-Instruct-FP8"
        self.data = None
        self.data_info = {}
        self.conversation_history = []
        
    def load_document(self, file_path: str) -> Dict[str, Any]:
        """
        Load and process different types of documents
        
        Args:
            file_path (str): Path to the document
            
        Returns:
            Dict containing processed data and metadata
        """
        file_extension = file_path.lower().split('.')[-1]
        
        try:
            if file_extension in ['csv']:
                return self._load_csv(file_path)
            elif file_extension in ['xlsx', 'xls']:
                return self._load_excel(file_path)
            elif file_extension == 'pdf':
                return self._load_pdf(file_path)
            elif file_extension in ['doc', 'docx']:
                return self._load_word(file_path)
            elif file_extension == 'txt':
                return self._load_text(file_path)
            elif file_extension in ['png', 'jpg', 'jpeg', 'tiff', 'bmp']:
                return self._load_image(file_path)
            else:
                return {"error": f"Unsupported file format: {file_extension}"}
                
        except Exception as e:
            return {"error": f"Error loading file: {str(e)}"}
    
    def _load_csv(self, file_path: str) -> Dict[str, Any]:
        """Load CSV file"""
        try:
            df = pd.read_csv(file_path)
            self.data = df
            self.data_info = {
                'type': 'tabular',
                'shape': df.shape,
                'columns': df.columns.tolist(),
                'dtypes': df.dtypes.astype(str).to_dict(),
                'missing_values': df.isnull().sum().to_dict(),
                'file_type': 'csv'
            }
            return {
                'success': True,
                'data_preview': df.head().to_dict(),
                'info': self.data_info
            }
        except Exception as e:
            return {"error": f"Error loading CSV: {str(e)}"}
    
    def _load_excel(self, file_path: str) -> Dict[str, Any]:
        """Load Excel file"""
        try:
            df = pd.read_excel(file_path)
            self.data = df
            self.data_info = {
                'type': 'tabular',
                'shape': df.shape,
                'columns': df.columns.tolist(),
                'dtypes': df.dtypes.astype(str).to_dict(),
                'missing_values': df.isnull().sum().to_dict(),
                'file_type': 'excel'
            }
            return {
                'success': True,
                'data_preview': df.head().to_dict(),
                'info': self.data_info
            }
        except Exception as e:
            return {"error": f"Error loading Excel: {str(e)}"}
    
    def _load_pdf(self, file_path: str) -> Dict[str, Any]:
        """Extract text from PDF"""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text()
            
            self.data = text
            self.data_info = {
                'type': 'text',
                'length': len(text),
                'pages': len(pdf_reader.pages),
                'file_type': 'pdf'
            }
            return {
                'success': True,
                'text_preview': text[:1000] + "..." if len(text) > 1000 else text,
                'info': self.data_info
            }
        except Exception as e:
            return {"error": f"Error loading PDF: {str(e)}"}
    
    def _load_word(self, file_path: str) -> Dict[str, Any]:
        """Extract text from Word document"""
        try:
            doc = docx.Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            self.data = text
            self.data_info = {
                'type': 'text',
                'length': len(text),
                'paragraphs': len(doc.paragraphs),
                'file_type': 'word'
            }
            return {
                'success': True,
                'text_preview': text[:1000] + "..." if len(text) > 1000 else text,
                'info': self.data_info
            }
        except Exception as e:
            return {"error": f"Error loading Word document: {str(e)}"}
    
    def _load_text(self, file_path: str) -> Dict[str, Any]:
        """Load text file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
            
            self.data = text
            self.data_info = {
                'type': 'text',
                'length': len(text),
                'lines': len(text.split('\n')),
                'file_type': 'text'
            }
            return {
                'success': True,
                'text_preview': text[:1000] + "..." if len(text) > 1000 else text,
                'info': self.data_info
            }
        except Exception as e:
            return {"error": f"Error loading text file: {str(e)}"}
    
    def _load_image(self, file_path: str) -> Dict[str, Any]:
        """Extract text from image using OCR"""
        try:
            image = Image.open(file_path)
            text = pytesseract.image_to_string(image)
            
            self.data = text
            self.data_info = {
                'type': 'image_text',
                'image_size': image.size,
                'text_length': len(text),
                'file_type': 'image'
            }
            return {
                'success': True,
                'extracted_text': text[:1000] + "..." if len(text) > 1000 else text,
                'info': self.data_info
            }
        except Exception as e:
            return {"error": f"Error processing image: {str(e)}"}
    
    def analyze_data(self) -> Dict[str, Any]:
        """
        Perform comprehensive data analysis
        """
        if self.data is None:
            return {"error": "No data loaded"}
        
        analysis = {}
        
        if self.data_info['type'] == 'tabular':
            analysis = self._analyze_tabular_data()
        elif self.data_info['type'] in ['text', 'image_text']:
            analysis = self._analyze_text_data()
        
        return analysis
    
    def _analyze_tabular_data(self) -> Dict[str, Any]:
        """Analyze tabular data (CSV, Excel)"""
        df = self.data
        
        analysis = {
            'basic_info': {
                'shape': df.shape,
                'columns': df.columns.tolist(),
                'data_types': df.dtypes.astype(str).to_dict(),
                'memory_usage': df.memory_usage(deep=True).sum()
            },
            'missing_data': {
                'missing_counts': df.isnull().sum().to_dict(),
                'missing_percentages': (df.isnull().sum() / len(df) * 100).to_dict()
            },
            'statistical_summary': {},
            'categorical_analysis': {},
            'correlations': {}
        }
        
        # Statistical summary for numerical columns
        numeric_cols = df.select_dtypes(include=[np.number]).columns
        if len(numeric_cols) > 0:
            analysis['statistical_summary'] = df[numeric_cols].describe().to_dict()
        
        # Categorical analysis
        categorical_cols = df.select_dtypes(include=['object', 'category']).columns
        for col in categorical_cols:
            if df[col].nunique() < 50:  # Only for columns with reasonable number of unique values
                analysis['categorical_analysis'][col] = {
                    'unique_values': df[col].nunique(),
                    'value_counts': df[col].value_counts().head(10).to_dict()
                }
        
        # Correlation analysis
        if len(numeric_cols) > 1:
            analysis['correlations'] = df[numeric_cols].corr().to_dict()
        
        return analysis
    
    def _analyze_text_data(self) -> Dict[str, Any]:
        """Analyze text data"""
        text = self.data
        
        analysis = {
            'basic_info': {
                'character_count': len(text),
                'word_count': len(text.split()),
                'line_count': len(text.split('\n')),
                'paragraph_count': len([p for p in text.split('\n\n') if p.strip()])
            },
            'word_frequency': {},
            'common_phrases': [] # Not implemented in original, keeping for structure
        }
        
        # Word frequency analysis
        words = re.findall(r'\b\w+\b', text.lower())
        word_freq = {}
        for word in words:
            if len(word) > 3:  # Filter out short words
                word_freq[word] = word_freq.get(word, 0) + 1
        
        # Top 20 most common words
        analysis['word_frequency'] = dict(sorted(word_freq.items(), key=lambda x: x[1], reverse=True)[:20])
        
        return analysis
    
    def create_visualization(self, viz_type: str, **kwargs) -> Optional[str]:
        """
        Create visualizations based on the data
        
        Args:
            viz_type (str): Type of visualization
            **kwargs: Additional parameters for visualization
            
        Returns:
            Base64 encoded image or Plotly JSON
        """
        if self.data is None or self.data_info['type'] != 'tabular':
            return None
        
        df = self.data
        
        try:
            if viz_type == 'correlation_heatmap':
                return self._create_correlation_heatmap(df)
            elif viz_type == 'distribution':
                return self._create_distribution_plot(df, kwargs.get('column'))
            elif viz_type == 'scatter':
                return self._create_scatter_plot(df, kwargs.get('x'), kwargs.get('y'))
            elif viz_type == 'bar':
                return self._create_bar_plot(df, kwargs.get('column'))
            elif viz_type == 'line':
                return self._create_line_plot(df, kwargs.get('x'), kwargs.get('y'))
            elif viz_type == 'box':
                return self._create_box_plot(df, kwargs.get('column'))
            else:
                return None
        except Exception as e:
            st.error(f"Error creating visualization: {str(e)}") # Using st.error for visibility
            return None
    
    def _create_correlation_heatmap(self, df: pd.DataFrame) -> str:
        """Create correlation heatmap"""
        numeric_cols = df.select_dtypes(include=[np.number]).columns
        if len(numeric_cols) < 2:
            st.warning("Not enough numeric columns for a correlation heatmap.")
            return None
        
        plt.figure(figsize=(10, 8))
        correlation_matrix = df[numeric_cols].corr()
        sns.heatmap(correlation_matrix, annot=True, cmap='coolwarm', center=0)
        plt.title('Correlation Heatmap')
        plt.tight_layout()
        
        # Convert to base64
        buffer = io.BytesIO()
        plt.savefig(buffer, format='png', dpi=300, bbox_inches='tight')
        buffer.seek(0)
        image_base64 = base64.b64encode(buffer.getvalue()).decode()
        plt.close()
        
        return image_base64
    
    def _create_distribution_plot(self, df: pd.DataFrame, column: str) -> str:
        """Create distribution plot"""
        if column not in df.columns:
            st.warning(f"Column '{column}' not found for distribution plot.")
            return None
        
        plt.figure(figsize=(10, 6))
        if df[column].dtype in ['int64', 'float64']:
            plt.hist(df[column].dropna(), bins=30, alpha=0.7, edgecolor='black')
            plt.xlabel(column)
            plt.ylabel('Frequency')
            plt.title(f'Distribution of {column}')
        else:
            value_counts = df[column].value_counts().head(10)
            plt.bar(range(len(value_counts)), value_counts.values)
            plt.xticks(range(len(value_counts)), value_counts.index, rotation=45)
            plt.xlabel(column)
            plt.ylabel('Count')
            plt.title(f'Top 10 Values in {column}')
        
        plt.tight_layout()
        
        # Convert to base64
        buffer = io.BytesIO()
        plt.savefig(buffer, format='png', dpi=300, bbox_inches='tight')
        buffer.seek(0)
        image_base64 = base64.b64encode(buffer.getvalue()).decode()
        plt.close()
        
        return image_base64
    
    def _create_scatter_plot(self, df: pd.DataFrame, x_col: str, y_col: str) -> str:
        """Create scatter plot"""
        if x_col not in df.columns or y_col not in df.columns:
            st.warning(f"Columns '{x_col}' or '{y_col}' not found for scatter plot.")
            return None
        if not pd.api.types.is_numeric_dtype(df[x_col]) or not pd.api.types.is_numeric_dtype(df[y_col]):
            st.warning(f"Both '{x_col}' and '{y_col}' must be numeric for a scatter plot.")
            return None
        
        plt.figure(figsize=(10, 6))
        plt.scatter(df[x_col], df[y_col], alpha=0.6)
        plt.xlabel(x_col)
        plt.ylabel(y_col)
        plt.title(f'{y_col} vs {x_col}')
        plt.tight_layout()
        
        # Convert to base64
        buffer = io.BytesIO()
        plt.savefig(buffer, format='png', dpi=300, bbox_inches='tight')
        buffer.seek(0)
        image_base64 = base64.b64encode(buffer.getvalue()).decode()
        plt.close()
        
        return image_base64
    
    def _create_bar_plot(self, df: pd.DataFrame, column: str) -> str:
        """Create bar plot"""
        if column not in df.columns:
            st.warning(f"Column '{column}' not found for bar plot.")
            return None
        
        plt.figure(figsize=(12, 6))
        value_counts = df[column].value_counts().head(15)
        if value_counts.empty:
            st.warning(f"Column '{column}' has no data to plot.")
            return None
        
        plt.bar(range(len(value_counts)), value_counts.values)
        plt.xticks(range(len(value_counts)), value_counts.index, rotation=45, ha='right')
        plt.xlabel(column)
        plt.ylabel('Count')
        plt.title(f'Top {min(len(value_counts), 15)} Values in {column}')
        plt.tight_layout()
        
        # Convert to base64
        buffer = io.BytesIO()
        plt.savefig(buffer, format='png', dpi=300, bbox_inches='tight')
        buffer.seek(0)
        image_base64 = base64.b64encode(buffer.getvalue()).decode()
        plt.close()
        
        return image_base64
    
    def _create_line_plot(self, df: pd.DataFrame, x_col: str, y_col: str) -> str:
        """Create line plot"""
        if x_col not in df.columns or y_col not in df.columns:
            st.warning(f"Columns '{x_col}' or '{y_col}' not found for line plot.")
            return None
        if not pd.api.types.is_numeric_dtype(df[y_col]):
            st.warning(f"Column '{y_col}' must be numeric for a line plot.")
            return None
        
        plt.figure(figsize=(12, 6))
        df_sorted = df.sort_values(x_col)
        plt.plot(df_sorted[x_col], df_sorted[y_col], marker='o', markersize=3)
        plt.xlabel(x_col)
        plt.ylabel(y_col)
        plt.title(f'{y_col} over {x_col}')
        plt.xticks(rotation=45)
        plt.tight_layout()
        
        # Convert to base64
        buffer = io.BytesIO()
        plt.savefig(buffer, format='png', dpi=300, bbox_inches='tight')
        buffer.seek(0)
        image_base64 = base64.b64encode(buffer.getvalue()).decode()
        plt.close()
        
        return image_base64
    
    def _create_box_plot(self, df: pd.DataFrame, column: str) -> str:
        """Create box plot"""
        if column not in df.columns or df[column].dtype not in ['int64', 'float64']:
            st.warning(f"Column '{column}' not found or is not numeric for box plot.")
            return None
        
        plt.figure(figsize=(8, 6))
        plt.boxplot(df[column].dropna())
        plt.ylabel(column)
        plt.title(f'Box Plot of {column}')
        plt.tight_layout()
        
        # Convert to base64
        buffer = io.BytesIO()
        plt.savefig(buffer, format='png', dpi=300, bbox_inches='tight')
        buffer.seek(0)
        image_base64 = base64.b64encode(buffer.getvalue()).decode()
        plt.close()
        
        return image_base64
    
    def ask_question(self, question: str) -> str:
        """
        Answer questions about the loaded data using the LLM
        
        Args:
            question (str): User's question about the data
            
        Returns:
            str: AI-generated answer
        """
        if self.data is None:
            return "No data has been loaded. Please load a document first."
        
        # Prepare context based on data type
        context = self._prepare_context_for_llm()
        
        # Create the prompt
        prompt = self._create_prompt(question, context)
        
        try:
            response = self.client.chat.completions.create(
                model=self.model,
                messages=[
                    {"role": "system", "content": "You are an expert data analyst. Provide detailed, accurate, and actionable insights based on the data provided. Use specific numbers and examples from the data when possible."},
                    {"role": "user", "content": prompt}
                ],
                max_tokens=1000,
                temperature=0.7
            )
            
            answer = response.choices[0].message.content
            
            # Store in conversation history - this is managed by Streamlit app state
            # self.conversation_history.append({
            #     'question': question,
            #     'answer': answer
            # })
            
            return answer
            
        except Exception as e:
            return f"Error generating response: {str(e)}"
    
    def _prepare_context_for_llm(self) -> str:
        """Prepare context information for the LLM"""
        context = f"Data Type: {self.data_info.get('type', 'N/A')}\n"
        context += f"File Type: {self.data_info.get('file_type', 'N/A')}\n\n"
        
        if self.data_info.get('type') == 'tabular':
            df = self.data
            context += f"Dataset Shape: {df.shape[0]} rows, {df.shape[1]} columns\n"
            context += f"Columns: {', '.join(df.columns.tolist())}\n\n"
            
            # Add sample data
            context += "Sample Data (first 5 rows):\n"
            context += df.head().to_string() + "\n\n"
            
            # Add statistical summary
            numeric_cols = df.select_dtypes(include=[np.number]).columns
            if len(numeric_cols) > 0:
                context += "Statistical Summary:\n"
                context += df[numeric_cols].describe().to_string() + "\n\n"
            
            # Add missing values info
            missing_info = df.isnull().sum()
            if missing_info.sum() > 0:
                context += "Missing Values:\n"
                for col, missing_count in missing_info.items():
                    if missing_count > 0:
                        context += f"{col}: {missing_count} ({missing_count/len(df)*100:.1f}%)\n"
                context += "\n"
        
        elif self.data_info.get('type') in ['text', 'image_text']:
            context += f"Text Length: {self.data_info.get('length', 'N/A')} characters\n"
            context += f"Sample Text:\n{self.data[:1000]}...\n\n"
        
        return context
    
    def _create_prompt(self, question: str, context: str) -> str:
        """Create a comprehensive prompt for the LLM"""
        prompt = f"""
I have loaded and analyzed a dataset. Here's the information about the data:

{context}

Based on this data, please answer the following question:
{question}

Please provide a detailed analysis that includes:
1. Direct answer to the question
2. Relevant insights from the data
3. Any patterns or trends you notice
4. Recommendations or next steps if applicable
5. Specific examples using actual data values when possible

If the question requires calculations or statistical analysis, please show your reasoning and provide specific numbers from the data.
"""
        return prompt
    
    def generate_insights(self) -> str:
        """
        Generate automatic insights about the loaded data
        """
        if self.data is None:
            return "No data has been loaded."
        
        context = self._prepare_context_for_llm()
        
        prompt = f"""
I have a dataset with the following information:

{context}

As an expert data analyst, please provide comprehensive insights about this data including:

1. **Data Overview**: Key characteristics and structure
2. **Key Findings**: Most important patterns, trends, or observations
3. **Data Quality**: Assessment of missing values, outliers, or data issues
4. **Statistical Insights**: Important statistical measures and what they mean
5. **Relationships**: Correlations or relationships between variables (if applicable)
6. **Actionable Recommendations**: What actions could be taken based on this data
7. **Potential Questions**: Interesting questions that could be explored further

Please be specific and use actual values from the data in your analysis.
"""
        
        try:
            response = self.client.chat.completions.create(
                model=self.model,
                messages=[
                    {"role": "system", "content": "You are an expert data analyst. Provide comprehensive, insightful, and actionable analysis of datasets. Use specific numbers and examples from the data."},
                    {"role": "user", "content": prompt}
                ],
                max_tokens=1500,
                temperature=0.7
            )
            
            return response.choices[0].message.content
            
        except Exception as e:
            return f"Error generating insights: {str(e)}"
    
    def suggest_visualizations(self) -> List[Dict[str, Any]]:
        """
        Suggest appropriate visualizations based on the data
        """
        if self.data is None or self.data_info.get('type') != 'tabular':
            return []
        
        df = self.data
        suggestions = []
        
        numeric_cols = df.select_dtypes(include=[np.number]).columns.tolist()
        categorical_cols = df.select_dtypes(include=['object', 'category']).columns.tolist()
        
        # Correlation heatmap
        if len(numeric_cols) > 1:
            suggestions.append({
                'type': 'correlation_heatmap',
                'title': 'Correlation Heatmap',
                'description': 'Shows relationships between numerical variables',
                'parameters': {}
            })
        
        # Distribution plots for numeric columns
        for col in numeric_cols[:3]:  # Limit to first 3 numeric columns
            suggestions.append({
                'type': 'distribution',
                'title': f'Distribution of {col}',
                'description': f'Shows the distribution of values in {col}',
                'parameters': {'column': col}
            })
        
        # Bar plots for categorical columns
        for col in categorical_cols[:3]:  # Limit to first 3 categorical columns
            if df[col].nunique() > 1 and df[col].nunique() <= 20:  # Only for columns with reasonable number of categories
                suggestions.append({
                    'type': 'bar',
                    'title': f'Value Counts for {col}',
                    'description': f'Shows the frequency of different values in {col}',
                    'parameters': {'column': col}
                })
        
        # Scatter plots for numeric pairs
        if len(numeric_cols) >= 2:
            suggestions.append({
                'type': 'scatter',
                'title': f'{numeric_cols[1]} vs {numeric_cols[0]}',
                'description': f'Shows relationship between {numeric_cols[0]} and {numeric_cols[1]}',
                'parameters': {'x': numeric_cols[0], 'y': numeric_cols[1]}
            })
        
        # Box plots for numeric columns
        for col in numeric_cols[:2]:  # Limit to first 2 numeric columns
            suggestions.append({
                'type': 'box',
                'title': f'Box Plot of {col}',
                'description': f'Shows distribution and outliers in {col}',
                'parameters': {'column': col}
            })
        
        # Line plots (if dates/time series or orderable numerical data)
        # This part is more heuristic, might need specific column names or types
        # For simplicity, will just suggest if two numeric columns are available
        if len(numeric_cols) >= 2:
             suggestions.append({
                'type': 'line',
                'title': f'Line Plot: {numeric_cols[1]} over {numeric_cols[0]}',
                'description': f'Shows trend of {numeric_cols[1]} values ordered by {numeric_cols[0]}',
                'parameters': {'x': numeric_cols[0], 'y': numeric_cols[1]}
            })

        return suggestions


# --- Streamlit App ---

st.set_page_config(
    page_title="Data Analyst AI Agent",
    page_icon="📊",
    layout="wide",
    initial_sidebar_state="expanded"
)

# --- Initialize Session State ---
if 'agent' not in st.session_state:
    st.session_state.agent = None
if 'data_loaded' not in st.session_state:
    st.session_state.data_loaded = False
if 'data_info' not in st.session_state:
    st.session_state.data_info = {}
if 'data_preview' not in st.session_state:
    st.session_state.data_preview = None
if 'chat_history' not in st.session_state:
    st.session_state.chat_history = []
if 'analysis_results' not in st.session_state:
    st.session_state.analysis_results = None
if 'current_tab' not in st.session_state:
    st.session_state.current_tab = "Overview"


# --- Sidebar for Configuration and Upload ---
with st.sidebar:
    st.title("⚙️ Configuration")
    
    st.markdown("---")
    st.subheader("Together.ai API Key")
    # Try to get API key from secrets.toml first
    api_key = "a911532c3f2b6e75cebc8868993637e359f3c1aca0206caa6c16818ce7dc73b5"
    if not api_key:
        api_key = st.text_input("Enter your Together.ai API Key:", type="password")
    
    if api_key and st.session_state.agent is None:
        try:
            st.session_state.agent = DataAnalystAgent(api_key)
            st.success("Agent Initialized! 🎉")
        except Exception as e:
            st.error(f"Error initializing agent: {e}. Please check your API key.")
            st.session_state.agent = None
    elif not api_key:
        st.info("Please enter your Together.ai API Key to initialize the agent.")

    st.markdown("---")
    st.subheader("Upload Document")
    uploaded_file = st.file_uploader(
        "Choose a file",
        type=['csv', 'xlsx', 'xls', 'pdf', 'doc', 'docx', 'txt', 'png', 'jpg', 'jpeg', 'tiff', 'bmp'],
        help="Supported formats: CSV, Excel, PDF, Word, TXT, Image (for OCR)"
    )

    if uploaded_file and st.session_state.agent:
        file_extension = uploaded_file.name.split('.')[-1].lower()
        
        # Save uploaded file to a temporary location
        with tempfile.NamedTemporaryFile(delete=False, suffix=f".{file_extension}") as tmp_file:
            tmp_file.write(uploaded_file.read())
            tmp_path = tmp_file.name

        with st.spinner(f"Loading {file_extension.upper()} file..."):
            load_result = st.session_state.agent.load_document(tmp_path)
            
            # Clean up the temporary file
            os.unlink(tmp_path)

            if load_result.get('success'):
                st.session_state.data_loaded = True
                st.session_state.data_info = load_result['info']
                if st.session_state.data_info['type'] == 'tabular':
                    st.session_state.data_preview = pd.DataFrame.from_dict(load_result['data_preview'])
                else:
                    st.session_state.data_preview = load_result['text_preview'] if 'text_preview' in load_result else load_result['extracted_text']
                
                st.success(f"File loaded successfully! Type: {st.session_state.data_info['type']}")
                st.session_state.chat_history = [] # Clear chat on new data load
                st.session_state.analysis_results = None # Clear analysis on new data load
                st.session_state.current_tab = "Overview" # Switch to overview tab
                st.rerun() # Rerun to update main content
            else:
                st.session_state.data_loaded = False
                st.error(f"Failed to load file: {load_result.get('error', 'Unknown error')}")
                st.session_state.agent.data = None
                st.session_state.data_info = {}
                st.session_state.data_preview = None
    elif uploaded_file and not st.session_state.agent:
        st.warning("Please initialize the agent with an API key first.")
    
    st.markdown("---")
    st.markdown("### About")
    st.info("This app uses a Data Analyst AI Agent powered by Together.ai's Llama-4-Maverick to analyze various document types and provide insights.")


# --- Main Content Area ---
st.title("📊 Data Analyst AI Agent")
st.markdown("Welcome to your AI-powered data analyst! Upload a document to get started.")

if not st.session_state.agent:
    st.warning("Please enter your Together.ai API Key in the sidebar to begin.")
elif not st.session_state.data_loaded:
    st.info("Upload a document (CSV, Excel, PDF, Word, TXT, Image) from the sidebar to analyze it.")

if st.session_state.data_loaded:
    tab_overview, tab_analysis, tab_insights, tab_qa, tab_viz = st.tabs(
        ["📄 Data Overview", "🔬 Data Analysis", "💡 Generated Insights", "💬 Ask a Question", "📈 Visualizations"]
    )

    with tab_overview:
        st.header("Document Information")
        if st.session_state.data_info:
            col1, col2 = st.columns(2)
            with col1:
                st.subheader("File Details")
                st.json(st.session_state.data_info)
            with col2:
                st.subheader("Data Preview")
                if st.session_state.data_info['type'] == 'tabular':
                    st.dataframe(st.session_state.data_preview)
                else:
                    st.text_area("Text Content Preview", st.session_state.data_preview, height=300)
        else:
            st.info("No data information available. Please load a file.")

    with tab_analysis:
        st.header("Comprehensive Data Analysis")
        st.markdown("Click the button below to perform a detailed analysis of the loaded data.")
        
        if st.button("Run Data Analysis", type="primary", use_container_width=True):
            if st.session_state.agent.data is None:
                st.warning("No data loaded to analyze.")
            else:
                with st.spinner("Analyzing data... This might take a moment."):
                    st.session_state.analysis_results = st.session_state.agent.analyze_data()
                    
                if st.session_state.analysis_results and "error" not in st.session_state.analysis_results:
                    st.success("Analysis complete!")
                else:
                    st.error(f"Analysis failed: {st.session_state.analysis_results.get('error', 'Unknown error')}")

        if st.session_state.analysis_results:
            analysis = st.session_state.analysis_results
            if "error" not in analysis:
                st.subheader("Analysis Results:")

                # Basic Info
                with st.expander("📊 Basic Information"):
                    if 'basic_info' in analysis:
                        st.json(analysis['basic_info'])

                # Missing Data
                with st.expander("🔍 Missing Data Overview"):
                    if 'missing_data' in analysis:
                        st.markdown("#### Missing Counts:")
                        st.json(analysis['missing_data'].get('missing_counts', {}))
                        st.markdown("#### Missing Percentages:")
                        st.json(analysis['missing_data'].get('missing_percentages', {}))

                # Statistical Summary
                if st.session_state.data_info['type'] == 'tabular':
                    with st.expander("📈 Statistical Summary (Numerical Columns)"):
                        if 'statistical_summary' in analysis and analysis['statistical_summary']:
                            # Convert to DataFrame for better display
                            summary_df = pd.DataFrame(analysis['statistical_summary'])
                            st.dataframe(summary_df)
                        else:
                            st.info("No numerical columns for statistical summary.")

                    # Categorical Analysis
                    with st.expander("🏷️ Categorical Analysis"):
                        if 'categorical_analysis' in analysis and analysis['categorical_analysis']:
                            for col, data in analysis['categorical_analysis'].items():
                                st.markdown(f"##### Column: `{col}`")
                                st.write(f"Unique Values: {data.get('unique_values')}")
                                st.markdown("Value Counts (Top 10):")
                                st.json(data.get('value_counts'))
                                st.markdown("---")
                        else:
                            st.info("No suitable categorical columns for analysis.")
                    
                    # Correlations
                    with st.expander("🔗 Correlations (Numerical Columns)"):
                        if 'correlations' in analysis and analysis['correlations']:
                            # Convert to DataFrame for better display
                            corr_df = pd.DataFrame(analysis['correlations'])
                            st.dataframe(corr_df)
                            st.info("A higher absolute value indicates a stronger correlation.")
                        else:
                            st.info("Not enough numerical columns for correlation analysis.")
                
                elif st.session_state.data_info['type'] in ['text', 'image_text']:
                    with st.expander("📝 Text Analysis"):
                        if 'word_frequency' in analysis:
                            st.markdown("#### Top 20 Word Frequencies:")
                            freq_df = pd.DataFrame(analysis['word_frequency'].items(), columns=['Word', 'Frequency'])
                            st.dataframe(freq_df)
                        else:
                            st.info("Text analysis results not available.")
            else:
                st.error(f"Error during analysis: {analysis['error']}")


    with tab_insights:
        st.header("Automated Insights")
        st.markdown("Get a comprehensive summary and actionable insights from the loaded data.")

        if st.button("Generate Insights", type="primary", use_container_width=True):
            if st.session_state.agent.data is None:
                st.warning("No data loaded to generate insights.")
            else:
                with st.spinner("Generating insights... This may take a minute."):
                    insights = st.session_state.agent.generate_insights()
                
                if "Error generating insights" in insights:
                    st.error(insights)
                else:
                    st.success("Insights Generated!")
                    st.session_state.generated_insights = insights
        
        if 'generated_insights' in st.session_state and st.session_state.generated_insights:
            st.markdown(st.session_state.generated_insights)


    with tab_qa:
        st.header("Ask the Data Analyst AI")
        st.markdown("Chat with the AI to get answers about your data.")

        # Display chat messages from history
        for message in st.session_state.chat_history:
            with st.chat_message(message["role"]):
                st.markdown(message["content"])

        # Chat input
        if prompt := st.chat_input("Ask a question about the data..."):
            st.session_state.chat_history.append({"role": "user", "content": prompt})
            with st.chat_message("user"):
                st.markdown(prompt)

            if st.session_state.agent.data is None:
                with st.chat_message("assistant"):
                    st.markdown("No data has been loaded. Please load a document first.")
                st.session_state.chat_history.append({"role": "assistant", "content": "No data has been loaded. Please load a document first."})
            else:
                with st.spinner("Thinking..."):
                    response = st.session_state.agent.ask_question(prompt)
                
                with st.chat_message("assistant"):
                    st.markdown(response)
                st.session_state.chat_history.append({"role": "assistant", "content": response})

    with tab_viz:
        st.header("Data Visualizations")

        if st.session_state.data_info['type'] != 'tabular':
            st.warning("Visualizations are only supported for tabular data (CSV, Excel).")
        elif not st.session_state.agent.data.empty:
            df_cols = st.session_state.agent.data.columns.tolist()
            numeric_cols = st.session_state.agent.data.select_dtypes(include=[np.number]).columns.tolist()
            
            st.subheader("Suggested Visualizations:")
            suggestions = st.session_state.agent.suggest_visualizations()
            if suggestions:
                cols = st.columns(3) # Display suggestions in up to 3 columns
                for i, viz_sugg in enumerate(suggestions):
                    with cols[i % 3]:
                        if st.button(f"Generate: {viz_sugg['title']}", key=f"sugg_viz_{i}", use_container_width=True):
                            with st.spinner(f"Generating {viz_sugg['title']}..."):
                                img_base64 = st.session_state.agent.create_visualization(
                                    viz_sugg['type'], **viz_sugg['parameters']
                                )
                                if img_base64:
                                    st.image(f"data:image/png;base64,{img_base64}", caption=viz_sugg['title'])
                                else:
                                    st.error(f"Could not generate {viz_sugg['title']}.")
            else:
                st.info("No specific visualization suggestions available for this data.")

            st.markdown("---")
            st.subheader("Custom Visualization Builder:")
            st.markdown("Select a plot type and relevant columns to create your own visualization.")

            plot_type = st.selectbox(
                "Select Plot Type",
                ["None", "distribution", "bar", "scatter", "line", "box", "correlation_heatmap"],
                key="custom_plot_type"
            )

            plot_params = {}
            if plot_type == "distribution":
                column = st.selectbox("Select Column", df_cols, key="dist_col")
                if column: plot_params['column'] = column
            elif plot_type == "bar":
                column = st.selectbox("Select Column", df_cols, key="bar_col")
                if column: plot_params['column'] = column
            elif plot_type == "scatter":
                col1, col2 = st.columns(2)
                with col1:
                    x_col = st.selectbox("Select X-axis Column (Numeric)", numeric_cols, key="scatter_x")
                with col2:
                    y_col = st.selectbox("Select Y-axis Column (Numeric)", numeric_cols, key="scatter_y")
                if x_col and y_col: plot_params.update({'x': x_col, 'y': y_col})
            elif plot_type == "line":
                col1, col2 = st.columns(2)
                with col1:
                    x_col = st.selectbox("Select X-axis Column", df_cols, key="line_x")
                with col2:
                    y_col = st.selectbox("Select Y-axis Column (Numeric)", numeric_cols, key="line_y")
                if x_col and y_col: plot_params.update({'x': x_col, 'y': y_col})
            elif plot_type == "box":
                column = st.selectbox("Select Column (Numeric)", numeric_cols, key="box_col")
                if column: plot_params['column'] = column
            
            if plot_type != "None" and st.button(f"Generate {plot_type.replace('_', ' ').title()} Plot", type="secondary", use_container_width=True):
                if (plot_type == "correlation_heatmap" and len(numeric_cols) < 2) or \
                   (plot_type in ["distribution", "bar", "box"] and 'column' not in plot_params) or \
                   (plot_type in ["scatter", "line"] and ('x' not in plot_params or 'y' not in plot_params)):
                    st.warning("Please select all required parameters for the chosen plot type.")
                else:
                    with st.spinner(f"Generating {plot_type.replace('_', ' ')} plot..."):
                        img_base64 = st.session_state.agent.create_visualization(plot_type, **plot_params)
                        if img_base64:
                            st.image(f"data:image/png;base64,{img_base64}", caption=f"{plot_type.replace('_', ' ').title()} Plot")
                        else:
                            st.error(f"Could not generate {plot_type.replace('_', ' ')} plot. Check parameters and data suitability.")
        else:
            st.info("Upload tabular data to enable visualization features.")
